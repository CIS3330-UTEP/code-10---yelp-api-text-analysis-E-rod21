from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
from yelpapi import YelpAPI
from docx import Document 

api_key = "WcUP7fRnOMxSjMMiW-GQnFZ8XzWGqkHG_TypojQopMMq6Y_gGq7v3Dt-2AS7X72ZVcWvrb4PrLXFFLt-9K6X9Wy6rZBBz83pAAlOI3-rvhDthDPl42CKBhHUCm9MZXYx"
yelp_api = YelpAPI(api_key)

search_term = "Mexican Restaurants"
location_term = "El Paso, TX"

search_result = yelp_api.search_query(
    term= search_term, location= location_term,
    sort_by = 'rating', limit = 10

)

# Sentiment Analysis 
def sentiment_analysis(reviews):
    analyze = SentimentIntensityAnalyzer()
    sentiments = []

    for review in reviews:
        text = review.get('text','')
        sentiment_score = analyze.polarity_scores(text)['compound']

        if sentiment_score >= .05:
            sentiment = 'Positive'
        elif sentiment_score <= -.05:
            sentiment = 'Negative'
        else:
            sentiment = "Neutral"

        sentiments.append(sentiment)

    return sentiments

doc = Document()

doc.add_heading("Yelp API Text Analysis", level = 0 )

doc.add_heading("Purpose: ", level=1)
doc.add_paragraph("In this report we want to find the 5 restaurants in El Paso, Tx that serve Mexican food. We also want the average amount spent at theses restuarants as well as 3 reviews per Restaurant as well as how the customers feel.")
doc.add_heading("Parameters:", level= 1)
doc.add_paragraph('The parameters for this search are; term = Mexican Restaurants, location = El Paso, TX, sorted by = Rating ')

for business in search_result['businesses']:

# Restaurant Name
    info = business
    print(info['name'])
    doc.add_heading(f"Restaurant: {info['name']}", level=1)

# Average Cost
    price_result = business.get('price', 'Not Available')
    print(price_result)
    doc.add_heading(f"Average Cost: {price_result}", level = 3)
    


    reviews = yelp_api.reviews_query(id=info['id'])

# Review
    for review in reviews['reviews'][:3]:
        review_text = review.get('text', '')
        print(review['text'])

        doc.add_heading("Review:", level=2)
        doc.add_paragraph(review_text)
    

# Sentiment 
    sentiments = sentiment_analysis(reviews['reviews'])
    print(sentiments)
    doc.add_paragraph(f"Sentiment: {sentiments}")

    doc.add_paragraph("\n")

doc.add_paragraph("Using the average cost we can see what the average spending is for the specific restaurant. Yelp uses $ as a representation of how much customers spend $ being the least and every additional represents higher costs. ")
doc.add_paragraph("The sentiment is calculated by the vaderSentiment, the idea behind it is that it should given an idea of how the customer felt about dinning/eating at said restaurant. It is calculated on Positive, Negative, or Neutral. However i ran into some issues with how vaderSentiment qualifies a review, this is because of a minor issues related to the customers post that might be positive but the way they expressed themselves in their review the code found it to be a Negative review. ")

doc.save('Yelp Analysis.docx')



